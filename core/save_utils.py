import re
from docx import Document
import os
from fpdf import FPDF
from core.scorer import section_relevance

def split_sections(cv_text, section_titles=None):
    """
    Split CV text into sections using regex based on section titles.
    Returns a list of (section_title, section_text).
    """
    if section_titles is None:
        section_titles = ['summary', 'experience', 'skills', 'education', 'projects', 'certifications', 'languages', 'interests']
    # Build regex pattern for section headers
    pattern = r'(^|\n)(' + '|'.join([re.escape(title) for title in section_titles]) + r')[:\n]'
    matches = list(re.finditer(pattern, cv_text, re.IGNORECASE))
    sections = []
    for i, match in enumerate(matches):
        start = match.end()
        end = matches[i+1].start() if i+1 < len(matches) else len(cv_text)
        title = match.group(2).strip()
        text = cv_text[start:end].strip()
        sections.append((title, text))
    if not sections:
        # fallback: treat all as one section
        sections = [("full", cv_text.strip())]
    return sections

def save_cv_to_docx(cv_text, output_path, section_titles=None, jd_text=None, highlight_relevance=False):
    """
    Save the tailored CV text to a DOCX file, preserving custom sections and optionally highlighting relevance.
    If highlight_relevance is True and jd_text is provided, most/least relevant sections are bolded/italicized.
    """
    doc = Document()
    sections = split_sections(cv_text, section_titles)
    top_sections, bottom_sections = ([], [])
    if highlight_relevance and jd_text:
        section_texts = [s[1] for s in sections]
        top, bottom = section_relevance(section_texts, jd_text, top_n=1)
        top_sections = [t[0] for t in top]
        bottom_sections = [b[0] for b in bottom]
    for title, text in sections:
        doc.add_heading(title, level=1)
        p = doc.add_paragraph(text)
        if highlight_relevance and jd_text:
            if text in top_sections:
                p.bold = True
            elif text in bottom_sections:
                p.italic = True
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc.save(output_path)

def save_cv_to_pdf(cv_text, output_path, section_titles=None, jd_text=None, highlight_relevance=False):
    """
    Save the tailored CV text to a PDF file, preserving custom sections and optionally highlighting relevance.
    If highlight_relevance is True and jd_text is provided, most/least relevant sections are bolded/italicized.
    """
    pdf = FPDF()
    pdf.add_page()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.set_font("Arial", size=12)
    sections = split_sections(cv_text, section_titles)
    top_sections, bottom_sections = ([], [])
    if highlight_relevance and jd_text:
        section_texts = [s[1] for s in sections]
        top, bottom = section_relevance(section_texts, jd_text, top_n=1)
        top_sections = [t[0] for t in top]
        bottom_sections = [b[0] for b in bottom]
    for title, text in sections:
        pdf.set_font("Arial", 'B', 14)
        pdf.cell(0, 10, title, ln=True)
        if highlight_relevance and jd_text:
            if text in top_sections:
                pdf.set_font("Arial", 'B', 12)
            elif text in bottom_sections:
                pdf.set_font("Arial", 'I', 12)
            else:
                pdf.set_font("Arial", size=12)
        else:
            pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, text)
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    pdf.output(output_path)
